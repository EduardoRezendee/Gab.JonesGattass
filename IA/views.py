from django.shortcuts import render, redirect, get_object_or_404
from .forms import PromptConfigForm, UploadBaseConhecimentoForm, UploadModelosForm, UploadDocumentoProcessoForm
from .models import PromptConfig, BaseConhecimento, ModelosDecisoes, DocumentoProcesso
from django.http import JsonResponse
from decouple import config
import logging
import openai
from django.views.decorators.csrf import csrf_exempt
import PyPDF2
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.core.exceptions import ValidationError
from docx import Document
import os 

# Configuração da API OpenAI
openai_client = openai.OpenAI(api_key=config("OPENAI_API_KEY", default=""))

# Configurar logs
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constantes
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

def editar_prompt(request):
    """Permite editar o prompt fixo da IA"""
    prompt_config = get_object_or_404(PromptConfig, pk=1)  # Supondo que o ID 1 seja o padrão

    if request.method == "POST":
        form = PromptConfigForm(request.POST, instance=prompt_config)
        if form.is_valid():
            form.save()
            return redirect("editar_prompt")
    else:
        form = PromptConfigForm(instance=prompt_config)

    return render(request, "editar_prompt.html", {"form": form})

def assistente_juridico(request):
    """ Exibe a interface do Assistente Jurídico """
    return render(request, "assistente_juridico.html")

@csrf_exempt
def ask_assistente_juridico(request):
    """Recebe mensagens do usuário e responde via OpenAI, gerando uma liminar de busca e apreensão baseada em modelos e jurisprudências."""
    if request.method == "POST":
        try:
            data = request.POST
            user_message = data.get("message", "").strip()

            if not user_message and "files" not in request.FILES:
                return JsonResponse({"response": "Por favor, digite uma pergunta ou envie um documento."}, status=400)

            # 1️⃣ **Recuperar o PromptConfig**
            prompt_config = PromptConfig.objects.first()
            if not prompt_config:
                prompt_config = PromptConfig.objects.create()

            prompt_texto = prompt_config.texto.strip()

            # 2️⃣ **Consultar a Base de Conhecimento (Jurisprudências)**
            jurisprudencias = BaseConhecimento.objects.exclude(arquivo="")

            # 3️⃣ **Consultar Modelos de Decisões (Usar como Template)**
            modelos = ModelosDecisoes.objects.all()
            modelo_template = "⚠️ Nenhum modelo de decisão válido encontrado."

            for modelo in modelos:
                file_extension = os.path.splitext(modelo.arquivo.name)[1].lower()

                try:
                    if file_extension == ".pdf":
                        modelo_template = extract_text_from_pdf(modelo.arquivo)
                    elif file_extension == ".docx":
                        modelo_template = extract_text_from_docx(modelo.arquivo)
                    else:
                        logger.warning(f"⚠️ Formato do modelo '{modelo.nome}' não suportado: {file_extension}")
                        continue

                    if modelo_template.strip():  # Se encontrar um modelo válido, usa ele e para a busca
                        break
                except Exception as e:
                    logger.warning(f"⚠️ Erro ao processar modelo '{modelo.nome}': {str(e)}")

            # 4️⃣ **Processar Documentos do Usuário**
            file_texts = []
            if "files" in request.FILES:
                for uploaded_file in request.FILES.getlist("files"):
                    file_extension = os.path.splitext(uploaded_file.name)[1].lower()

                    if file_extension == ".pdf":
                        file_texts.append(extract_text_from_pdf(uploaded_file))
                    elif file_extension == ".docx":
                        file_texts.append(extract_text_from_docx(uploaded_file))
                    else:
                        file_texts.append(f"⚠️ Formato não suportado ({file_extension}): {uploaded_file.name}")

            file_text = "\n".join(file_texts) if file_texts else "Nenhum documento enviado."

            # 5️⃣ **Extrair textos da Base de Conhecimento (Jurisprudências)**
            jurisprudencia_textos = []
            for j in jurisprudencias:
                file_extension = os.path.splitext(j.arquivo.name)[1].lower()
                try:
                    if file_extension == ".pdf":
                        jurisprudencia_textos.append(extract_text_from_pdf(j.arquivo))
                    elif file_extension == ".docx":
                        jurisprudencia_textos.append(extract_text_from_docx(j.arquivo))
                    else:
                        logger.warning(f"⚠️ Formato não suportado para jurisprudência: {file_extension}")
                except Exception as e:
                    logger.warning(f"⚠️ Erro ao processar jurisprudência '{j.nome}': {str(e)}")

            # 6️⃣ **Montar o Contexto para a IA**
            contexto = f"""
            **INSTRUÇÃO:**
            {prompt_texto}

            **MODELO DE DECISÃO (Template a ser seguido exatamente):**
            {modelo_template}

            **Jurisprudências Relevantes:**
            {"\n\n".join(jurisprudencia_textos) if jurisprudencia_textos else "Nenhuma jurisprudência encontrada."}

            **Documentos do Usuário:**
            {file_text}

            **Pergunta do Usuário:**
            {user_message}
            """

            # 7️⃣ **Chamada para a OpenAI**
            messages = [
                {"role": "system", "content": prompt_texto},
                {"role": "user", "content": contexto}
            ]

            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",  # Use o modelo adequado
                messages=messages
            )

            # 8️⃣ **Retornar a Liminar Gerada**
            bot_response = response.choices[0].message.content.strip()
            return JsonResponse({"response": bot_response})

        except openai.OpenAIError as e:  # ✅ Corrigindo a captura de erro da OpenAI
            logger.error(f"❌ Erro na API da OpenAI: {str(e)}", exc_info=True)
            return JsonResponse({"response": "⚠️ Erro na comunicação com a OpenAI. Tente novamente mais tarde."}, status=500)

        except Exception as e:
            logger.error(f"❌ Erro ao processar mensagem: {str(e)}", exc_info=True)
            return JsonResponse({"response": f"Erro ao processar a solicitação: {str(e)}"}, status=500)

    return JsonResponse({"error": "Método não permitido"}, status=405)


# Função para extrair texto de arquivos DOCX 📄
def extract_text_from_docx(file):
    try:
        doc = Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip() if text.strip() else "⚠️ O documento DOCX está vazio."
    except Exception as e:
        logger.error(f"❌ Erro ao processar DOCX: {str(e)}", exc_info=True)
        return "⚠️ Erro ao processar o documento DOCX."


def extract_text_from_pdf(file):
    try:
        reader = PyPDF2.PdfReader(file)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        return text.strip() if text else "⚠️ O PDF não contém texto extraível."
    except PyPDF2.errors.PdfReadError:
        return "⚠️ Erro ao processar o PDF: arquivo corrompido ou formato inválido."

# Função para extrair texto de DOCX 📄
def extract_text_from_docx(file):
    try:
        doc = Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip() if text else "⚠️ O documento DOCX está vazio."
    except Exception as e:
        logger.error(f"❌ Erro ao processar DOCX: {str(e)}", exc_info=True)
        return "⚠️ Erro ao processar o documento DOCX."

def upload_base_conhecimento(request):
    """Realiza upload de base de conhecimento."""
    if request.method == "POST":
        form = UploadBaseConhecimentoForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect("upload_base")  # Redireciona para evitar reenvio do formulário
    else:
        form = UploadBaseConhecimentoForm()

    # Busca todos os documentos já enviados
    documentos = BaseConhecimento.objects.all()

    # Passa o formulário e os documentos para o template
    return render(request, "upload_base.html", {"form": form, "documentos": documentos})

def upload_modelos(request):
    """Realiza upload de modelos e decisões."""
    if request.method == "POST":
        form = UploadModelosForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect("upload_modelos")  # Redireciona após salvar o documento

    else:
        form = UploadModelosForm()

    # Busca todos os documentos salvos
    documentos = ModelosDecisoes.objects.all()

    return render(request, "upload_modelos.html", {"form": form, "documentos": documentos})


def upload_documento_processo(request):
    """Realiza upload de documentos do processo."""
    if request.method == "POST":
        try:
            form = UploadDocumentoProcessoForm(request.POST, request.FILES)
            if form.is_valid():
                form.save()
                return JsonResponse({"message": "Documento enviado com sucesso!"}, status=200)
            else:
                return JsonResponse({"error": "Erro no formulário. Verifique os campos."}, status=400)
        except Exception as e:
            logger.error(f"Erro ao salvar documento: {e}", exc_info=True)
            return JsonResponse({"error": f"Erro interno: {str(e)}"}, status=500)

    # Exibe os documentos já enviados
    documentos = DocumentoProcesso.objects.all()
    return render(request, "upload_documento_processo.html", {"documentos": documentos})

def painel_ia(request):
    """Exibe a tela inicial do Painel IA"""
    return render(request, "painel_ia.html")

def deletar_base_conhecimento(request, doc_id):
    if request.method == "DELETE":
        try:
            documento = get_object_or_404(BaseConhecimento, id=doc_id)
            documento.delete()
            return JsonResponse({"status": "success", "message": "Documento deletado com sucesso."})
        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=500)
    return JsonResponse({"status": "error", "message": "Método não permitido."}, status=405)